import io
import re
import time
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
import streamlit as st
from docx import Document

# ---- Page config
st.set_page_config(page_title="PDF → DOCX (Hybrid OCR)", page_icon="📄", layout="centered")

st.title("📄 PDF → DOCX Converter (Hybrid OCR + Parser)")
st.caption("Digital PDFs use parsing; scanned/image pages use OCR (Tesseract).")

# ---- Sidebar options
with st.sidebar:
    st.header("Settings")
    mode = st.radio(
        "Processing mode",
        ["Auto (Hybrid)", "Force OCR (all pages)", "Parser only (no OCR)"],
        index=0,
        help="Hybrid tries text parsing first, then OCR if needed.",
    )
    zoom = st.slider("OCR render scale (quality vs speed)", 1.0, 3.0, 2.0, 0.25)
    ocr_lang = st.text_input("Tesseract language code(s)", "eng",
                             help="e.g. 'eng', 'eng+urd'. Add language packs in packages.txt.")
    add_page_headings = st.checkbox("Add 'Page N' headings in DOCX", value=True)

def clean_text(text: str) -> str:
    """Remove characters not valid in DOCX (XML) to avoid save errors."""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text or "")

def page_needs_ocr(parsed_text: str) -> bool:
    """Heuristic: if parsed text is empty or almost empty, fall back to OCR."""
    if not parsed_text:
        return True
    letters = sum(c.isalpha() for c in parsed_text)
    return letters < 25  # very little actual text ⇒ likely scanned

def pixmap_to_pil(pix: fitz.Pixmap) -> Image.Image:
    mode = "RGBA" if pix.alpha else "RGB"
    return Image.frombytes(mode, [pix.width, pix.height], pix.samples)

def convert_pdf(pdf_bytes: bytes, mode: str, zoom: float, ocr_lang: str, add_page_headings: bool) -> bytes:
    """
    Convert PDF bytes to a DOCX (bytes) using hybrid parsing + OCR.
    """
    # Ensure Tesseract path (typical on Streamlit Cloud with packages.txt)
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

    # Open once for PyMuPDF (rendering + parsing)
    doc_fitz = fitz.open(stream=pdf_bytes, filetype="pdf")
    # Open again for pdfplumber (table extraction on digital pages)
    doc_plumber = pdfplumber.open(io.BytesIO(pdf_bytes))

    out_doc = Document()

    progress = st.progress(0.0, text="Starting…")
    total = len(doc_fitz)

    for i in range(total):
        page_fitz = doc_fitz[i]
        page_plumber = doc_plumber.pages[i]

        # 1) Try text parsing (digital PDFs)
        parsed_text = ""
        if mode in ("Auto (Hybrid)", "Parser only (no OCR)"):
            # PyMuPDF text
            parsed_text = page_fitz.get_text("text")
            parsed_text = clean_text(parsed_text)

        # 2) Try table extraction (only for digital pages)
        tables = []
        if mode in ("Auto (Hybrid)", "Parser only (no OCR)"):
            try:
                tables = page_plumber.extract_tables()
            except Exception:
                tables = []

        # 3) Decide if we need OCR
        need_ocr = (mode == "Force OCR (all pages)") or (mode == "Auto (Hybrid)" and page_needs_ocr(parsed_text))

        page_text_written = False
        if add_page_headings:
            out_doc.add_heading(f"Page {i+1}", level=2)

        # 4) If we have useful parsed text, add it first
        if parsed_text and not need_ocr:
            out_doc.add_paragraph(parsed_text)
            page_text_written = True

        # 5) If there are tables, write them to the DOCX
        if tables and not need_ocr:
            for tbl in tables:
                if not tbl or not tbl[0]:
                    continue
                rows = len(tbl)
                cols = len(tbl[0])
                t = out_doc.add_table(rows=rows, cols=cols)
                for r in range(rows):
                    for c in range(cols):
                        val = tbl[r][c]
                        t.cell(r, c).text = clean_text(str(val) if val is not None else "")
                out_doc.add_paragraph("")  # spacing

        # 6) OCR fallback (or forced)
        if need_ocr:
            # Render page to image at higher zoom for better OCR
            mat = fitz.Matrix(zoom, zoom)
            pix = page_fitz.get_pixmap(matrix=mat, alpha=False)  # RGB
            pil_img = pixmap_to_pil(pix)
            ocr_text = pytesseract.image_to_string(pil_img, lang=ocr_lang)
            ocr_text = clean_text(ocr_text)
            out_doc.add_paragraph(ocr_text if ocr_text.strip() else "[No text detected]")
            page_text_written = True

        # 7) If nothing was written at all, add a placeholder
        if not page_text_written:
            out_doc.add_paragraph("[No extractable content on this page]")

        progress.progress((i+1)/total, text=f"Processed page {i+1}/{total}")
        time.sleep(0.02)  # tiny pause so the progress bar visibly updates

    # Save DOCX to bytes
    bio = io.BytesIO()
    out_doc.save(bio)
    bio.seek(0)
    doc_fitz.close()
    doc_plumber.close()
    return bio.read()

# ---- App body
uploaded = st.file_uploader("Upload a PDF", type=["pdf"])

if uploaded is not None:
    st.info("File received. Click **Convert** to start.", icon="✅")

    if st.button("Convert"):
        with st.spinner("Converting… this may take a moment for large PDFs"):
            pdf_bytes = uploaded.read()
            try:
                docx_bytes = convert_pdf(
                    pdf_bytes=pdf_bytes,
                    mode=mode,
                    zoom=zoom,
                    ocr_lang=ocr_lang,
                    add_page_headings=add_page_headings,
                )
                st.success("Done! Download your DOCX below.")
                st.download_button(
                    label="⬇️ Download DOCX",
                    data=docx_bytes,
                    file_name=(uploaded.name.rsplit(".", 1)[0] + ".docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"Conversion failed: {e}")
else:
    st.caption("Maximum upload size is set by Streamlit Cloud (typically up to ~200 MB).")
