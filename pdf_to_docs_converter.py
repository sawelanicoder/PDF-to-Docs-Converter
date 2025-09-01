!apt-get install -y tesseract-ocr
!pip install pymupdf pytesseract pillow opencv-python pandas numpy python-docx beautifulsoup4

import fitz  # PyMuPDF
import pytesseract
import cv2
import numpy as np
import pandas as pd
from PIL import Image
from docx import Document
from bs4 import BeautifulSoup
from google.colab import files
import os

def upload_pdf():
    print("📂 Please upload your PDF file...")
    uploaded = files.upload()
    for fn in uploaded.keys():
        print(f"✅ File '{fn}' uploaded successfully!")
        return fn  # return uploaded filename

pdf_path = upload_pdf()

def pdf_to_images(pdf_path, zoom=2):
    doc = fitz.open(pdf_path)
    images = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    return images

images = pdf_to_images(pdf_path)
print(f"📄 Extracted {len(images)} pages as images.")

def ocr_images(images):
    text_pages = []
    for i, img in enumerate(images):
        text = pytesseract.image_to_string(img)
        text_pages.append(text)
        print(f"🔍 OCR done for page {i+1}")
    return text_pages

ocr_text = ocr_images(images)

import re

def clean_text(text):
    """Removes non-XML compatible characters from text."""
    # Remove characters that are not valid in XML 1.0
    # This includes control characters but excludes valid ones like newline (\n) and tab (\t)
    cleaned_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return cleaned_text

# Save as plain text
with open("output_text.txt", "w", encoding="utf-8") as f:
    for page_num, page_text in enumerate(ocr_text, start=1):
        f.write(f"--- Page {page_num} ---\n")
        f.write(page_text + "\n\n")

# Save as Word doc
doc = Document()
for page_num, page_text in enumerate(ocr_text, start=1):
    cleaned_page_text = clean_text(page_text) # Clean the text
    doc.add_heading(f"Page {page_num}", level=2)
    doc.add_paragraph(cleaned_page_text)
doc.save("output_text.docx")

print("✅ OCR results saved as 'output_text.txt' and 'output_text.docx'")

from google.colab import files
files.download("output_text.docx")

